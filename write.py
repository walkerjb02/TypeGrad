import webbrowser as wb
from docx import Document
import os

def writetxt(text):
    with open("TypeGradOutput.txt", "w") as file:
        file.write(f"TypeGrad Output\n\n{text}")
        fal = os.path.abspath("TypeGradOutput.txt")
        wb.open(fal)

def write_doc(text):
    document = Document()
    document.add_heading("TypeGrad Output", 0)
    document.add_paragraph(f"{text}")
    document.save("TypeGradResult.docx")
    fal = os.path.abspath('TypeGradResult.docx')
    wb.open(fal)
